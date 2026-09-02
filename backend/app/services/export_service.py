from __future__ import annotations

import json
import logging
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

logger = logging.getLogger(__name__)

# Standalone image line: ![alt](url) - the format produced by finalize_figures.
_IMAGE_LINE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)\s]+)\)\s*$")
# Inline images anywhere in the text (markdown export path rewriting).
_INLINE_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)\s]+)\)")
# **bold** emphasis segments (figure captions like **图1：** ...).
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def ensure_export_dir(base_dir: Path, project_id: str) -> Path:
    from app.services.storage import get_storage_backend
    path_str = get_storage_backend().ensure_export_dir(project_id)
    return Path(path_str) if not path_str.startswith("s3://") else Path(path_str)


def _resolve_image_file(url: str) -> Path | None:
    """Map an in-content image URL to a file on disk.

    Workflow figures carry API paths (``/api/projects/{pid}/images/...``)
    which live under ``backend/data/storage``. Unresolvable URLs return None
    so callers can fall back to writing the original text.
    """
    if not url:
        return None
    if url.startswith("/api/projects/"):
        from app.database import backend_dir

        rel = url[len("/api/projects/"):].split("?", 1)[0]
        candidate = backend_dir / "data" / "storage" / rel
        return candidate if candidate.exists() else None
    local = Path(url)
    return local if local.exists() else None


def export_markdown(target_dir: Path, filename: str, content_md: str) -> Path:
    """Write markdown and bundle referenced images next to it.

    Body images point at runtime API URLs (``/api/projects/...``) which break
    the moment the backend is offline. Export copies each image into
    ``images/`` and rewrites the markdown to relative paths so the export is
    self-contained and shareable.
    """
    images_dir = target_dir / "images"
    used_names: set[str] = set()

    def _rewrite(match: re.Match[str]) -> str:
        alt, url = match.group(1), match.group(2)
        src = _resolve_image_file(url)
        if src is None:
            logger.warning("export_markdown: image not found on disk, left as-is: %s", url)
            return match.group(0)
        images_dir.mkdir(parents=True, exist_ok=True)
        dest_name = src.name
        counter = 1
        while dest_name in used_names:
            dest_name = f"{src.stem}-{counter}{src.suffix}"
            counter += 1
        used_names.add(dest_name)
        shutil.copy2(src, images_dir / dest_name)
        return f"![{alt}](images/{dest_name})"

    path = target_dir / filename
    path.write_text(_INLINE_IMAGE_RE.sub(_rewrite, content_md), encoding="utf-8")
    return path


def _add_docx_paragraph_with_emphasis(doc: Any, text: str) -> None:
    """Add a paragraph rendering **bold** segments as bold runs."""
    pos = 0
    paragraph = None
    for match in _BOLD_RE.finditer(text):
        before = text[pos:match.start()]
        if before and paragraph is None:
            paragraph = doc.add_paragraph(before)
        elif before and paragraph is not None:
            paragraph.add_run(before)
        if paragraph is None:
            paragraph = doc.add_paragraph()
        paragraph.add_run(match.group(1)).bold = True
        pos = match.end()
    if pos < len(text):
        rest = text[pos:]
        if paragraph is None:
            doc.add_paragraph(rest)
        else:
            paragraph.add_run(rest)


def export_docx(target_dir: Path, filename: str, content_md: str) -> Path:
    path = target_dir / filename
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore

        doc = Document()
        for raw_line in content_md.splitlines():
            line = raw_line.strip()
            image_match = _IMAGE_LINE_RE.match(line)
            if image_match:
                src = _resolve_image_file(image_match.group(2))
                if src is not None:
                    try:
                        doc.add_picture(str(src), width=Inches(5.8))
                        continue
                    except Exception as exc:
                        logger.warning("export_docx: failed to embed %s (%s), writing as text", src, exc)
                # Unresolvable image: fall through and keep the markdown text.
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line:
                _add_docx_paragraph_with_emphasis(doc, line)
        doc.save(path)
        return path
    except Exception:
        # Graceful fallback to plain text with .docx suffix.
        path.write_text(content_md, encoding="utf-8")
        return path


def export_pdf(target_dir: Path, filename: str, content_md: str) -> Path:
    path = target_dir / filename
    try:
        from fpdf import FPDF  # type: ignore

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        # Register CJK font for Chinese support
        font_path = Path(__file__).resolve().parent.parent.parent / "fonts" / "simhei.ttf"
        if font_path.exists():
            pdf.add_font("SimHei", "", str(font_path), uni=True)
            pdf.set_font("SimHei", size=11)
        else:
            pdf.set_font("Helvetica", size=11)
        for raw_line in content_md.splitlines():
            line = raw_line.strip() or " "
            image_match = _IMAGE_LINE_RE.match(line)
            if image_match:
                src = _resolve_image_file(image_match.group(2))
                if src is not None:
                    try:
                        pdf.ln(2)
                        pdf.image(str(src), w=pdf.epw)
                        pdf.ln(2)
                        continue
                    except Exception as exc:
                        logger.warning("export_pdf: failed to embed %s (%s), writing as text", src, exc)
            # fpdf renders no markdown: strip ** emphasis so captions like
            # **图1：** do not show literal asterisks.
            text = _BOLD_RE.sub(r"\1", line)
            pdf.multi_cell(0, 7, txt=text)
        pdf.output(str(path))
        return path
    except Exception:
        # Final fallback: write plain text bytes to preserve endpoint availability.
        path.write_bytes(content_md.encode("utf-8", errors="ignore"))
        return path


def export_bibtex(target_dir: Path, filename: str, papers: Iterable[dict]) -> Path:
    from app.services.citation_service import format_bibliography

    class _FakePaper:
        def __init__(self, data: dict) -> None:
            self.id = data.get("id", "")
            self.title = data.get("title") or "Untitled"
            self.authors = data.get("authors") or []
            self.year = data.get("year")
            self.venue = data.get("venue") or ""
            self.doi = data.get("doi") or ""
            self.arxiv_id = data.get("arxiv_id") or ""
            self.source_url = data.get("source_url") or ""
            self.pdf_url = data.get("pdf_url") or ""

    fake_papers = [_FakePaper(p) for p in papers]
    text = format_bibliography(fake_papers, style="bibtex")
    path = target_dir / filename
    path.write_text(text + "\n", encoding="utf-8")
    return path


def export_json(target_dir: Path, filename: str, payload: dict | list) -> Path:
    path = target_dir / filename
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def export_quality_report(
    target_dir: Path,
    filename: str,
    *,
    draft_version: int,
    review_rounds: list[dict[str, Any]],
    final_metrics: dict[str, Any],
    publication_prepared: bool,
) -> Path:
    """Export a structured quality gate report alongside the draft."""
    report = {
        "report_type": "quality_gate",
        "draft_version": draft_version,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "publication_prepared": publication_prepared,
        "final_metrics": final_metrics,
        "review_history": review_rounds,
        "thresholds": {
            "evidence_coverage": 0.90,
            "citation_validity": 0.90,
            "logic_score": 0.80,
            "style_score": 0.80,
            "critical_issues": 0,
            "unsupported_claims": 0,
        },
    }
    return export_json(target_dir, filename, report)

